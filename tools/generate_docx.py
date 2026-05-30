#!/usr/bin/env python3
"""Convert a CV markdown file to a styled Word .docx resume.

Usage: generate_docx.py [cv.md] [output.docx]
Defaults: cv.md -> output/resume-<slugified-name>.docx

Supports: # name (title), ## section, ### role/subhead, - bullets,
**bold** inline, and plain paragraphs. Pure python-docx, no network.
"""
import os, re, sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x2A, 0x44)
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_runs(paragraph, text):
    """Render inline **bold** segments."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build(md_path, out_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            h = doc.add_heading(level=0)
            r = h.add_run(line[2:].strip())
            r.font.color.rgb = ACCENT
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("### "):
            p = doc.add_paragraph()
            add_runs(p, line[4:].strip())
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = ACCENT
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line.lstrip()[2:].strip())
        else:
            p = doc.add_paragraph()
            add_runs(p, line.strip())

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    md = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, "cv.md")
    if len(sys.argv) > 2:
        out = sys.argv[2]
    else:
        name = ""
        for l in open(md, encoding="utf-8"):
            if l.startswith("# "):
                name = l[2:].strip()
                break
        slug = re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-") or "resume"
        out = os.path.join(ROOT, "output", f"resume-{slug}.docx")
    print("wrote", build(md, out))


if __name__ == "__main__":
    main()
